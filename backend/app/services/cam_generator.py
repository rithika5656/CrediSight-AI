"""
CAM Report Generator
Generates structured Credit Appraisal Memo as PDF and Word documents.
"""
import os
import io
from datetime import datetime, timezone
from typing import Optional

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable,
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


PRIMARY_COLOR = HexColor("#1a365d")
ACCENT_COLOR = HexColor("#2b6cb0")
LIGHT_BG = HexColor("#f7fafc")


def generate_cam_pdf(application_data: dict, output_path: str) -> str:
    """Generate a PDF Credit Appraisal Memo."""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="CamTitle", fontSize=20, leading=24, textColor=PRIMARY_COLOR,
        alignment=TA_CENTER, spaceAfter=12, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="SectionHeader", fontSize=14, leading=18, textColor=PRIMARY_COLOR,
        spaceAfter=8, spaceBefore=16, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="SubHeader", fontSize=11, leading=14, textColor=ACCENT_COLOR,
        spaceAfter=6, spaceBefore=8, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="CamBody", fontSize=10, leading=14, spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="CamFooter", fontSize=8, leading=10, textColor=HexColor("#718096"),
        alignment=TA_CENTER,
    ))

    elements = []

    # Title
    elements.append(Paragraph("CREDIT APPRAISAL MEMO", styles["CamTitle"]))
    elements.append(Paragraph("CrediSight AI -- Corporate Credit Decisioning Platform", styles["CamFooter"]))
    elements.append(Spacer(1, 0.3 * inch))
    elements.append(HRFlowable(width="100%", thickness=2, color=PRIMARY_COLOR))
    elements.append(Spacer(1, 0.2 * inch))

    # Application Info
    app_info = application_data.get("application", {})
    elements.append(Paragraph("1. APPLICATION SUMMARY", styles["SectionHeader"]))
    info_data = [
        ["Application ID", str(app_info.get("id", "N/A"))],
        ["Company Name", app_info.get("company_name", "N/A")],
        ["CIN Number", app_info.get("cin_number", "N/A")],
        ["GST Number", app_info.get("gst_number", "N/A")],
        ["Industry Sector", app_info.get("industry_sector", "N/A")],
        ["Requested Loan Amount", f"INR {app_info.get('requested_loan_amount', 0):,.2f}"],
        ["Application Date", app_info.get("created_at", "N/A")],
    ]
    table = Table(info_data, colWidths=[2.5 * inch, 4 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), LIGHT_BG),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#e2e8f0")),
        ("PADDING", (0, 0), (-1, -1), 8),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 0.2 * inch))

    # Company Overview
    elements.append(Paragraph("2. COMPANY OVERVIEW", styles["SectionHeader"]))
    desc = app_info.get("business_description", "No description provided.")
    elements.append(Paragraph(desc, styles["CamBody"]))
    elements.append(Spacer(1, 0.15 * inch))

    # Financial Analysis
    elements.append(Paragraph("3. FINANCIAL ANALYSIS", styles["SectionHeader"]))
    financial = application_data.get("financial_data", {})

    gst = financial.get("gst_filing", {})
    if gst:
        elements.append(Paragraph("GST Filing Analysis", styles["SubHeader"]))
        elements.append(Paragraph(f"Reported Revenue: INR {gst.get('reported_revenue', 0):,.2f}", styles["CamBody"]))
        elements.append(Paragraph(f"Total Transactions: {gst.get('total_transactions', 0)}", styles["CamBody"]))

    bank = financial.get("bank_statement", {})
    if bank:
        elements.append(Paragraph("Bank Statement Analysis", styles["SubHeader"]))
        elements.append(Paragraph(f"Total Credits: INR {bank.get('total_credits', 0):,.2f}", styles["CamBody"]))
        elements.append(Paragraph(f"Total Debits: INR {bank.get('total_debits', 0):,.2f}", styles["CamBody"]))
        elements.append(Paragraph(f"Average Balance: INR {bank.get('average_balance', 0):,.2f}", styles["CamBody"]))

    itr = financial.get("income_tax_return", {})
    if itr:
        elements.append(Paragraph("Income Tax Return Analysis", styles["SubHeader"]))
        elements.append(Paragraph(f"Gross Income: INR {itr.get('gross_income', 0):,.2f}", styles["CamBody"]))
        elements.append(Paragraph(f"Tax Paid: INR {itr.get('tax_paid', 0):,.2f}", styles["CamBody"]))
    elements.append(Spacer(1, 0.15 * inch))

    # Cross-Verification
    elements.append(Paragraph("4. FINANCIAL CROSS-VERIFICATION", styles["SectionHeader"]))
    cv = application_data.get("cross_verification", {})
    elements.append(Paragraph(f"Verification Score: {cv.get('verification_score', 'N/A')}/100", styles["CamBody"]))
    if cv.get("revenue_mismatch"):
        elements.append(Paragraph(
            f"Revenue Mismatch Detected: {cv.get('mismatch_percentage', 0)}%",
            styles["CamBody"],
        ))
    suspicious = cv.get("suspicious_patterns", [])
    for pattern in suspicious:
        elements.append(Paragraph(f"  - {pattern}", styles["CamBody"]))
    if not suspicious:
        elements.append(Paragraph("No suspicious patterns detected.", styles["CamBody"]))
    elements.append(Spacer(1, 0.15 * inch))

    # Risk Assessment
    elements.append(Paragraph("5. RISK ASSESSMENT", styles["SectionHeader"]))
    risk = application_data.get("risk_assessment", {})
    elements.append(Paragraph(f"Composite Risk Score: {risk.get('composite_score', 'N/A')}/100", styles["CamBody"]))
    elements.append(Paragraph(f"Risk Level: {risk.get('risk_level', 'N/A').upper()}", styles["CamBody"]))
    elements.append(Spacer(1, 0.1 * inch))

    # Five Cs
    elements.append(Paragraph("6. FIVE Cs EVALUATION", styles["SectionHeader"]))
    five_cs = risk.get("five_cs", {})
    for c_name, c_data in five_cs.items():
        score = c_data.get("score", 0)
        details = c_data.get("details", "")
        elements.append(Paragraph(f"{c_name.upper()} (Score: {score}/100)", styles["SubHeader"]))
        elements.append(Paragraph(details, styles["CamBody"]))
    elements.append(Spacer(1, 0.15 * inch))

    # Final Recommendation
    elements.append(Paragraph("7. FINAL LENDING RECOMMENDATION", styles["SectionHeader"]))
    rec = application_data.get("recommendation", {})
    decision = rec.get("decision", "N/A").upper()
    elements.append(Paragraph(f"Decision: {decision}", styles["SubHeader"]))
    elements.append(Paragraph(f"Recommended Loan Limit: INR {rec.get('recommended_loan_limit', 0):,.2f}", styles["CamBody"]))
    elements.append(Paragraph(f"Suggested Interest Rate: {rec.get('suggested_interest_rate', 0)}% p.a.", styles["CamBody"]))
    elements.append(Spacer(1, 0.1 * inch))

    reasoning = rec.get("reasoning", [])
    if reasoning:
        elements.append(Paragraph("Reasoning:", styles["SubHeader"]))
        for reason in reasoning:
            elements.append(Paragraph(f"  - {reason}", styles["CamBody"]))

    conditions = rec.get("conditions", [])
    if conditions:
        elements.append(Spacer(1, 0.1 * inch))
        elements.append(Paragraph("Conditions:", styles["SubHeader"]))
        for cond in conditions:
            elements.append(Paragraph(f"  - {cond}", styles["CamBody"]))

    # Footer
    elements.append(Spacer(1, 0.5 * inch))
    elements.append(HRFlowable(width="100%", thickness=1, color=HexColor("#cbd5e0")))
    elements.append(Spacer(1, 0.1 * inch))
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    elements.append(Paragraph(f"Generated by CrediSight AI on {now}", styles["CamFooter"]))
    elements.append(Paragraph("This document is system-generated and for internal use only.", styles["CamFooter"]))

    doc.build(elements)
    return output_path


def generate_cam_docx(application_data: dict, output_path: str) -> str:
    """Generate a Word document Credit Appraisal Memo."""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = DocxDocument()

    # Styles
    style = doc.styles["Title"]
    style.font.color.rgb = RGBColor(26, 54, 93)

    # Title
    title = doc.add_heading("CREDIT APPRAISAL MEMO", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("CrediSight AI -- Corporate Credit Decisioning Platform")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(113, 128, 150)
    subtitle.runs[0].font.size = Pt(10)

    doc.add_paragraph("")

    app_info = application_data.get("application", {})

    # 1. Application Summary
    doc.add_heading("1. APPLICATION SUMMARY", level=1)
    table = doc.add_table(rows=7, cols=2)
    table.style = "Light Grid Accent 1"
    fields = [
        ("Application ID", str(app_info.get("id", "N/A"))),
        ("Company Name", app_info.get("company_name", "N/A")),
        ("CIN Number", app_info.get("cin_number", "N/A")),
        ("GST Number", app_info.get("gst_number", "N/A")),
        ("Industry Sector", app_info.get("industry_sector", "N/A")),
        ("Requested Loan Amount", f"INR {app_info.get('requested_loan_amount', 0):,.2f}"),
        ("Application Date", app_info.get("created_at", "N/A")),
    ]
    for i, (label, value) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # 2. Company Overview
    doc.add_heading("2. COMPANY OVERVIEW", level=1)
    doc.add_paragraph(app_info.get("business_description", "No description provided."))

    # 3. Financial Analysis
    doc.add_heading("3. FINANCIAL ANALYSIS", level=1)
    financial = application_data.get("financial_data", {})
    gst = financial.get("gst_filing", {})
    if gst:
        doc.add_heading("GST Filing Analysis", level=2)
        doc.add_paragraph(f"Reported Revenue: INR {gst.get('reported_revenue', 0):,.2f}")
        doc.add_paragraph(f"Total Transactions: {gst.get('total_transactions', 0)}")

    bank = financial.get("bank_statement", {})
    if bank:
        doc.add_heading("Bank Statement Analysis", level=2)
        doc.add_paragraph(f"Total Credits: INR {bank.get('total_credits', 0):,.2f}")
        doc.add_paragraph(f"Total Debits: INR {bank.get('total_debits', 0):,.2f}")
        doc.add_paragraph(f"Average Balance: INR {bank.get('average_balance', 0):,.2f}")

    itr = financial.get("income_tax_return", {})
    if itr:
        doc.add_heading("Income Tax Return Analysis", level=2)
        doc.add_paragraph(f"Gross Income: INR {itr.get('gross_income', 0):,.2f}")
        doc.add_paragraph(f"Tax Paid: INR {itr.get('tax_paid', 0):,.2f}")

    # 4. Cross-Verification
    doc.add_heading("4. FINANCIAL CROSS-VERIFICATION", level=1)
    cv = application_data.get("cross_verification", {})
    doc.add_paragraph(f"Verification Score: {cv.get('verification_score', 'N/A')}/100")
    if cv.get("revenue_mismatch"):
        doc.add_paragraph(f"Revenue Mismatch: {cv.get('mismatch_percentage', 0)}%")
    for pattern in cv.get("suspicious_patterns", []):
        doc.add_paragraph(f"  - {pattern}")

    # 5. Risk Assessment
    doc.add_heading("5. RISK ASSESSMENT", level=1)
    risk = application_data.get("risk_assessment", {})
    doc.add_paragraph(f"Composite Risk Score: {risk.get('composite_score', 'N/A')}/100")
    doc.add_paragraph(f"Risk Level: {risk.get('risk_level', 'N/A').upper()}")

    # 6. Five Cs
    doc.add_heading("6. FIVE Cs EVALUATION", level=1)
    five_cs = risk.get("five_cs", {})
    for c_name, c_data in five_cs.items():
        doc.add_heading(f"{c_name.upper()} (Score: {c_data.get('score', 0)}/100)", level=2)
        doc.add_paragraph(c_data.get("details", ""))

    # 7. Recommendation
    doc.add_heading("7. FINAL LENDING RECOMMENDATION", level=1)
    rec = application_data.get("recommendation", {})
    p = doc.add_paragraph()
    run = p.add_run(f"Decision: {rec.get('decision', 'N/A').upper()}")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph(f"Recommended Loan Limit: INR {rec.get('recommended_loan_limit', 0):,.2f}")
    doc.add_paragraph(f"Suggested Interest Rate: {rec.get('suggested_interest_rate', 0)}% p.a.")

    for reason in rec.get("reasoning", []):
        doc.add_paragraph(f"  - {reason}")

    if rec.get("conditions"):
        doc.add_heading("Conditions", level=2)
        for cond in rec["conditions"]:
            doc.add_paragraph(f"  - {cond}")

    # Footer
    doc.add_paragraph("")
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    footer = doc.add_paragraph(f"Generated by CrediSight AI on {now}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(113, 128, 150)

    doc.save(output_path)
    return output_path
